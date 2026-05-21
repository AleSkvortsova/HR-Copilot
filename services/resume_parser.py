from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text_from_resume(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(file_path)
    if suffix in {".doc", ".docx"}:
        return _extract_from_docx(file_path)

    raise ValueError(f"Неподдерживаемый формат файла: {suffix}")


def _extract_from_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    result = "\n\n".join(parts).strip()
    if not result:
        raise ValueError("Не удалось извлечь текст из PDF. Возможно, файл содержит только изображения.")
    return result


def _extract_from_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    result = "\n\n".join(parts).strip()
    if not result:
        raise ValueError("Не удалось извлечь текст из докумена Word.")
    return result
